import requests
from bs4 import BeautifulSoup
import json
from docx import Document

# Load the URL from config.json
with open('config.json') as config_file:
    config = json.load(config_file)
    url = config['url']

response = requests.get(url)
soup = BeautifulSoup(response.text, 'html.parser')

elements = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'span', 'ul'])

# Create a Word document
doc = Document()
doc.add_heading('TFT Website Scraped Data', 0)

for element in elements:
    # Handle headers (h1-h6)
    if element.name.startswith('h'):
        level = int(element.name[1])
        doc.add_heading(element.text.strip(), level=level)
        next_sibling = element.find_next_sibling()
        while next_sibling and next_sibling.name == 'p':
            doc.add_paragraph(next_sibling.text.strip())
            next_sibling = next_sibling.find_next_sibling()

    elif element.name == 'span':
        doc.add_paragraph(element.text.strip())

    elif element.name == 'ul':
        for li in element.find_all('li'):
            doc.add_paragraph(f"• {li.text.strip()}", style='List Bullet')

doc.save('tft_scraped_data.docx')

print('Data has been saved to tft_scraped_data.docx')
