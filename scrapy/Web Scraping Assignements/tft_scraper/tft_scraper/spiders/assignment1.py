import scrapy
from docx import Document

class TFTSpider1(scrapy.Spider):
    name = 'assignment1'
    start_urls = ['https://www.tftus.com/']

    def parse(self, response):
        doc = Document()
        doc.add_heading('TFT Website Scraped Data', 0)

        headers = response.xpath('//h1 | //h2 | //h3 | //h4 | //h5 | //h6')
        for header in headers:
            level = int(header.root.tag[1:])
            header_texts = header.xpath('.//text()').getall()
            full_text = ' '.join(header_texts).strip()
            doc.add_heading(full_text, level=level)

            next_sibling = header.xpath('following-sibling::*[self::p]')
            for p in next_sibling:
                p_text = p.xpath('.//text()').getall()
                p_text = ' '.join(p_text).strip()
                if p_text:
                    doc.add_paragraph(p_text)

        for element in response.xpath('//p | //span | //div | //a'):
            element_texts = element.xpath('.//text()').getall()
            full_text = ' '.join(element_texts).strip()
            if full_text:
                if element.root.tag == 'a':
                    link = element.xpath('@href').get()
                    full_text = f"{full_text} ({link})"
                doc.add_paragraph(full_text)

        list_items = response.xpath('//ul/li')
        for li in list_items:
            li_text = li.xpath('.//text()').getall()
            li_text = ' '.join(li_text).strip()
            if li_text:
                doc.add_paragraph(f"• {li_text}", style='List Bullet')

        # Save the document
        doc.save('tft_spider1_data.docx')
        self.log('Word document saved as tft_spider1_data.docx')
